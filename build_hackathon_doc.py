from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = "AI陪伴机器人_项目提案_角色分工_赛前与48小时流程.docx"
BLUE = "275D73"
LIGHT = "EAF2F5"
PALE = "F4F7F8"
INK = "20343D"
GRAY = "5E6B72"


def set_font(run, size=10.5, bold=False, color=INK):
    run.font.name = "STHeiti"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "STHeiti")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "STHeiti")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "STHeiti")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color=INK, size=9.2, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell_margins(cell)


def set_table_widths(table, widths):
    table.autofit = False
    total = int(sum(widths) * 1440)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_w = row.cells[idx]._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for i, h in enumerate(headers):
        shade(table.rows[0].cells[i], BLUE)
        set_cell_text(table.rows[0].cells[i], h, bold=True, color="FFFFFF", size=9.2,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        if ridx % 2:
            for c in cells:
                shade(c, PALE)
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, bold=(i == 0), size=9.0,
                          align=WD_ALIGN_PARAGRAPH.LEFT if i else WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_font(r, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(text))
    return p


doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(0.7)
sec.left_margin = sec.right_margin = Inches(0.75)
sec.header_distance = sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "STHeiti"
normal._element.rPr.rFonts.set(qn("w:ascii"), "STHeiti")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "STHeiti")
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.15
for name, size, before, after in (("Heading 1", 16, 14, 7), ("Heading 2", 12.5, 10, 5)):
    st = styles[name]
    st.font.name = "STHeiti"
    st._element.rPr.rFonts.set(qn("w:ascii"), "STHeiti")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "STHeiti")
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(BLUE)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)

# Header / footer
hp = sec.header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(hp.add_run("硬件黑客松｜赛前准备 + 48 小时现场执行版 v0.2"), 8.5, color=GRAY)
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(fp.add_run("内部讨论稿 · 角色待认领 · 可按技术验证结果调整"), 8, color=GRAY)

# Title block
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(4)
set_font(p.add_run("桌面 AI 陪伴机器人"), 24, True, BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(14)
set_font(p.add_run("项目提案 · 角色分工 · 赛前准备与比赛流程"), 13, True, GRAY)

add_heading(doc, "一、Proposal｜项目提案")
add_heading(doc, "1. 一句话定位", 2)
add_body(doc, "一款能够判断陪伴时机、主动靠近用户，并通过表情、语音和动作提供轻量陪伴的桌面 AI 机器人。")

table = doc.add_table(rows=1, cols=1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_widths(table, [6.9])
shade(table.cell(0, 0), LIGHT)
set_cell_text(table.cell(0, 0), "不是等你开口，而是在合适的时候主动来到你身边。", bold=True,
              color=BLUE, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph().paragraph_format.space_after = Pt(0)

add_heading(doc, "2. 产品形态", 2)
add_table(doc, ["模块", "MVP 作用"], [
    ("T5 开发板/屏幕", "显示眼睛、表情与运行状态"),
    ("摄像头", "检测用户是否在场及大致方向"),
    ("麦克风与扬声器", "接收简单指令，播放主动问候与反馈"),
    ("双轮小车底盘", "完成转向、靠近、停止与离开"),
    ("3D 打印结构/软质外壳", "固定硬件并形成亲和、可替换的角色形象"),
], [2.0, 4.9])

add_heading(doc, "3. 目标用户", 2)
for x in [
    "经常独自学习、远程工作或长时间伏案的年轻人",
    "独居、初到陌生城市或日常陪伴较少的人",
    "压力较大，但不一定会主动寻求帮助的人",
    "需要轻量陪伴，同时重视安静、隐私和互动边界的人",
]: add_bullet(doc, x)
add_body(doc, "设计原则：产品不限定性别；女性友好体现在安全感、互动边界、细腻体验和非刻板化外观，而不是简单使用粉色元素。")

add_heading(doc, "4. 核心使用场景", 2)
add_table(doc, ["优先级", "场景", "机器人行为"], [
    ("Demo 主场景", "久坐学习/工作", "发现用户长时间伏案后主动靠近，询问是否需要休息或陪伴"),
    ("扩展", "独处回家", "发现用户出现后主动迎接并进行简短问候"),
    ("扩展", "情绪低落", "用低打扰方式询问，或保持安静陪伴"),
    ("扩展", "专注时刻", "识别勿扰状态，降低存在感并停止主动交互"),
], [1.15, 1.55, 4.2])

add_heading(doc, "5. 核心功能与范围", 2)
add_table(doc, ["级别", "功能", "验收标准"], [
    ("必须", "用户检测与方向判断", "能判断有人/无人，并输出大致方向"),
    ("必须", "主动移动", "能转向、靠近，在安全距离停止并可离开"),
    ("必须", "表情反馈", "至少实现眨眼、开心、关心、等待四种状态"),
    ("必须", "主动问候", "可播放预录音或 TTS，并支持接受/拒绝分支"),
    ("必须", "陪伴策略", "能根据触发条件决定靠近、互动或保持安静"),
    ("可选", "情绪识别/自由对话/记忆", "只在核心闭环稳定后增加"),
], [0.75, 2.2, 3.95])

add_heading(doc, "6. 核心创新")
for x in [
    "从被动问答转向能够发起互动的主动陪伴",
    "通过移动、距离、表情和声音创造实体存在感",
    "强调有分寸的主动：理解接受、拒绝、无回应和勿扰状态",
    "软件人格与实体外壳均可扩展，不替代真实人际关系",
]: add_bullet(doc, x)

add_heading(doc, "二、Division of Labor｜五类角色（待认领）")
add_body(doc, "当前只定义项目所需角色，不分配成员姓名。角色可一人多岗，也可多人协作；认领后由项目统筹统一记录。")
add_table(doc, ["角色", "主要职责", "阶段性交付物"], [
    ("项目统筹与系统集成", "控制范围与进度；统一软硬件接口；组织联调、测试和风险降级", "接口清单、进度看板、可运行整机、兜底方案"),
    ("机器人硬件与运动控制", "底盘、电机、供电、电路、摄像头安装及基础运动", "硬件连接图、运动控制指令、安全停止"),
    ("AI 感知与交互逻辑", "人物检测、状态触发、语音交互、陪伴状态机和行为决策", "感知输出、状态机、语音/触发接口"),
    ("产品体验与视觉设计", "产品定义、交互流程、表情动画、UI、外壳与整体视觉", "交互稿、表情素材、控制界面、外观方案"),
    ("Demo 展示与路演表达", "设计演示剧本；整合产品故事、PPT、视频及现场讲解", "Demo 脚本、路演材料、备份视频、讲解稿"),
], [1.6, 3.3, 2.0])

add_heading(doc, "角色协作规则", 2)
for x in [
    "每个模块必须先提供最小可用接口，再继续优化效果。",
    "接口统一使用少量明确指令，例如 STOP、MOVE、FACE_CARE、SPEAK_HELLO。",
    "任何新增功能不得影响核心 Demo 闭环；第 36 小时后原则上停止加功能。",
    "所有角色共同参与最终联调与至少 10 次完整演示测试。",
]: add_bullet(doc, x)

add_heading(doc, "三、核心 Demo 流程")
add_table(doc, ["步骤", "系统行为", "现场呈现"], [
    ("1 感知", "检测用户在场、方向及久坐触发条件", "机器人从休眠状态醒来并看向用户"),
    ("2 判断", "判断当前是否适合互动", "屏幕显示观察/关心表情"),
    ("3 靠近", "转向并移动，在预设安全距离停止", "机器人主动来到用户附近"),
    ("4 问候", "播放关心语音并等待用户回应", "“你已经忙很久了，需要休息一下吗？”"),
    ("5 分支", "接受则陪伴；拒绝或无回应则安静退出", "用不同表情、语音和动作体现边界感"),
], [0.8, 3.1, 3.0])

add_heading(doc, "四、赛前准备流程｜8 月 10 日晚—13 日")
add_table(doc, ["时间", "阶段目标", "关键动作", "完成标准"], [
    ("8 月 10 日晚", "团队对齐", "确认产品定位、主场景、五类角色和沟通方式", "形成共识文档，列出待确认事项"),
    ("8 月 11 日", "方案与物料", "确定 T5、摄像头、底盘、供电、扬声器与外壳方案", "完成物料清单，确认自带/采购/现场获取"),
    ("8 月 12 日", "技术预研", "分别验证屏幕、电机、视觉、语音和通信方式", "关键模块有最小测试结果和接口草案"),
    ("8 月 13 日", "赛前定稿", "确定 Demo 剧本、代码仓库、文件结构和比赛首日顺序", "设备带齐；任务可认领；风险与降级方案明确"),
], [1.1, 1.25, 3.0, 1.55])

add_heading(doc, "五、正式比赛流程｜8 月 14 日开始，共 48 小时")
add_table(doc, ["比赛时间", "阶段目标", "关键动作", "完成标准"], [
    ("0–4 h", "现场确认", "检查现场物料与规则；完成角色认领；锁定最终 MVP", "屏幕点亮、轮子转动、Demo 剧本确定"),
    ("4–16 h", "模块跑通", "硬件、感知、表情、语音、UI 和路演内容并行开发", "每个模块均可独立演示"),
    ("16–28 h", "首次联调", "先用按钮触发，打通移动—表情—语音闭环", "完成一次粗糙但完整的 Demo"),
    ("28–36 h", "自动化与体验", "接入人物/久坐触发，优化距离、外观和拒绝分支", "核心流程可重复运行"),
    ("36–44 h", "冻结与测试", "停止扩功能；修复问题；连续演示并录制备份", "连续成功至少 10 次"),
    ("44–48 h", "路演准备", "充电、固定接线、排练、准备手动与视频兜底", "90 秒 Demo 与完整讲解就绪"),
], [0.9, 1.15, 3.0, 1.85])

add_heading(doc, "六、降级与风险方案")
add_table(doc, ["理想能力", "降级方案"], [
    ("情绪识别", "改用久坐计时、特定动作或控制台按钮触发"),
    ("实时语音识别", "改用实体按钮、网页选项或关键词识别"),
    ("LLM 自由对话", "使用预设对话和分支，确保现场稳定"),
    ("精确跟随与导航", "限定演示区域，按人物方向移动固定时长"),
    ("在线 TTS", "提前准备本地录音"),
    ("完整 3D 外壳", "使用毛绒玩具、布套或快速打印简化结构"),
], [2.2, 4.7])

add_heading(doc, "七、团队确认清单")
for x in [
    "确认产品名称与角色形象",
    "确认开发板、摄像头、底盘、电源和扬声器是否可用",
    "确认五类角色的认领方式与系统集成人",
    "确认久坐/状态变化采用自动识别还是可控触发",
    "确认首次完整联调时间和第 36 小时功能冻结点",
    "确认 90 秒 Demo 剧本、手动控制和备份视频方案",
]: add_bullet(doc, "□ " + x)

doc.save(OUT)
print(OUT)
